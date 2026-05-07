import os
import sys
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--break-system-packages"])
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    base_dir = os.path.abspath(os.path.dirname(__file__))
    doc = Document()
    
    # Title
    title = doc.add_heading('Data Cleaning & Preprocessing Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Dataset: Diabetes 130-US Hospitals (1999-2008)').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Goal: Predict 30-Day Hospital Readmission').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Introduction
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph("This document outlines the comprehensive data cleaning and preprocessing blueprint "
                      "based on the original Strack et al. (2014) paper and modern Machine Learning literature "
                      "(2023-2024). These steps are mandatory for preparing the dataset for predictive modeling "
                      "(e.g., XGBoost, Random Forest, etc.) to achieve academic-grade results.")
    
    # Phase 1
    doc.add_heading('2. Phase 1: Record Filtering (The Strack Baseline)', level=1)
    p = doc.add_paragraph("These steps are necessary to avoid data leakage and logical errors:")
    
    ul1 = doc.add_paragraph(style='List Bullet')
    ul1.add_run('Isolate First Encounters: ').bold = True
    ul1.add_run('The dataset contains multiple visits for the same patients. We must keep only the first encounter per patient. '
                'Keeping subsequent visits causes "data leakage" because the model will memorize the patient instead of '
                'learning the disease patterns. (Reduces rows from 101,766 to ~71,518).')
                
    ul2 = doc.add_paragraph(style='List Bullet')
    ul2.add_run('Remove Terminal Discharges: ').bold = True
    ul2.add_run('Patients who died or were discharged to a hospice (Discharge IDs: 11, 13, 14, 19, 20, 21) cannot be readmitted. '
                'Including them confuses the model (it predicts "No Readmission" for a terminally ill patient). '
                '(Reduces rows to ~69,984).')
    
    # Phase 2
    doc.add_heading('3. Phase 2: Dimensionality Reduction & Feature Engineering', level=1)
    doc.add_paragraph("Modern algorithms struggle with hundreds of sparse categories. We must compress them.")
    
    ul3 = doc.add_paragraph(style='List Bullet')
    ul3.add_run('Target Variable Binarization: ').bold = True
    ul3.add_run("Change the readmitted column from 3 classes ('<30', '>30', 'NO') into 2 classes: 1 (Readmitted within 30 days) "
                "vs 0 (Others). This aligns with hospital penalty metrics.")
                
    ul4 = doc.add_paragraph(style='List Bullet')
    ul4.add_run('Drop Sparse Columns: ').bold = True
    ul4.add_run("Drop 'weight' (97% missing) and 'payer_code' (40% missing) entirely.")
    
    ul5 = doc.add_paragraph(style='List Bullet')
    ul5.add_run('Handle Missing Values (?): ').bold = True
    ul5.add_run("Replace '?' with 'Missing' in medical_specialty, and 'Unknown' in race.")
    
    ul6 = doc.add_paragraph(style='List Bullet')
    ul6.add_run('ICD-9 Diagnosis Grouping: ').bold = True
    ul6.add_run("Map the 900+ unique disease codes in diag_1, diag_2, and diag_3 into 9 broad clinical categories based on "
                "HCUP standards: Circulatory, Respiratory, Digestive, Diabetes, Injury, Musculoskeletal, Genitourinary, "
                "Neoplasms, and Other.")
                
    ul7 = doc.add_paragraph(style='List Bullet')
    ul7.add_run('Medication Compression: ').bold = True
    ul7.add_run("Drop extremely rare drugs. Create a new numerical feature 'num_medications_changed' by counting how many drugs "
                "have 'Up' or 'Down' for that patient. Frequent dosage changes indicate instability.")
                
    ul8 = doc.add_paragraph(style='List Bullet')
    ul8.add_run('Age Transformation: ').bold = True
    ul8.add_run("Convert string brackets '[40-50)' into ordinal integers or midpoints (e.g., 45) so the model understands "
                "the mathematical progression of age.")
                
    # Phase 3
    doc.add_heading('4. Phase 3: Model Preparation (2024 Best Practices)', level=1)
    
    ul9 = doc.add_paragraph(style='List Bullet')
    ul9.add_run('Scaling Numeric Features: ').bold = True
    ul9.add_run("Apply StandardScaler (Z-score normalization) to features like time_in_hospital, num_lab_procedures.")
    
    ul10 = doc.add_paragraph(style='List Bullet')
    ul10.add_run('Handling Class Imbalance: ').bold = True
    ul10.add_run("Only ~11% of the filtered patients are readmitted. Apply SMOTE or class weights in XGBoost to penalize "
                 "the model heavily when it misses a readmitted patient.")
                 
    # Future Plans
    doc.add_heading('5. Future Architecture: NLP & In-App Cleaning', level=1)
    doc.add_paragraph("In the upcoming phase, the data cleaning pipeline will be directly integrated into the React/Node JS "
                      "frontend application. Additionally, we will introduce Natural Language Processing (NLP) vectorization "
                      "(e.g., Word2Vec/TF-IDF) to categorize complex clinical free-text or uncategorized features automatically, "
                      "rather than relying strictly on hardcoded rules.")
                      
    doc.save(os.path.join(base_dir, 'Comprehensive_Data_Cleaning_Report.docx'))
    print("Docx created successfully.")

if __name__ == '__main__':
    create_report()
