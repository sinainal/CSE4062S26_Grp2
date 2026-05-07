#!/usr/bin/env python3
"""Generate a professional DOCX report describing the current app workflow."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE = Path(__file__).resolve().parent
OUT_DIR = BASE / "output" / "doc"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_PATH = OUT_DIR / "APP_Final_Workflow_Report.docx"


def load_json(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def set_document_defaults(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for style_name, size in [("Title", 22), ("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 11)]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        if style_name == "Title":
            style.font.bold = True

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)


def add_page_number_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CSE4062S26_Grp2 - App Workflow Report")
    run.font.size = Pt(8.5)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CSE4062S26_Grp2 App Final Workflow Report")
    run.bold = True
    run.font.size = Pt(22)

    for text in [
        "Medical Data Mining, Descriptive Analytics, and Predictive Analytics",
        "Dataset: Diabetes 130-US Hospitals (1999-2008)",
        "Prepared from the current dashboard-centered application state",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(11)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Date: 2026-05-07")
    r.italic = True

    doc.add_page_break()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered_list(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def main() -> None:
    baseline = load_json(BASE / "user_tools/visualisation_tool/baseline_model_report.json")
    regression = load_json(BASE / "user_tools/visualisation_tool/regression_report.json")
    model_lab = load_json(BASE / "user_tools/visualisation_tool/model_lab_report.json")
    feature_selection = load_json(BASE / "user_tools/visualisation_tool/feature_selection_report.json")
    feature_subset = load_json(BASE / "user_tools/visualisation_tool/feature_subset_report.json")
    clustering = load_json(BASE / "user_tools/visualisation_tool/clustering_report.json")
    apriori = load_json(BASE / "user_tools/visualisation_tool/association_rules_report.json")
    cleaning = load_json(BASE / "user_tools/visualisation_tool/cleaning_data.json")
    model_ready = load_json(BASE / "user_tools/visualisation_tool/model_ready_data.json")
    academic = load_json(BASE / "user_tools/visualisation_tool/academic_data.json")

    doc = Document()
    set_document_defaults(doc)
    add_page_number_footer(doc.sections[0])
    add_title_page(doc)

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "This report summarizes the final dashboard-oriented version of the project repository. "
        "The application is designed as a single-click local analytics environment that bootstraps its own Python "
        "environment, regenerates missing outputs when necessary, and opens a browser-based medical data mining dashboard."
    )
    doc.add_paragraph(
        "The current implementation covers the full project scope: data understanding and preprocessing, descriptive "
        "analytics through clustering and Apriori, predictive analytics through classification and regression, and a live "
        "model lab that allows repeated experiment runs with different hyperparameters and dataset sizes."
    )

    doc.add_heading("2. Current App Snapshot", level=1)
    add_table(
        doc,
        ["Item", "Current Value"],
        [
            ["Raw dataset rows", "101,766"],
            ["Raw dataset columns", "50"],
            ["Cleaned dataset rows", str(cleaning.get("stats", {}).get("kept_rows", 69_970))],
            ["Model-ready rows", str(model_ready.get("model_ready_rows", 69_970))],
            ["Model-ready columns", str(model_ready.get("model_ready_columns", 160))],
            ["Baseline classifiers", str(len(baseline.get("models", [])))],
            ["Live lab algorithms", str(len(model_lab.get("algorithms", [])))],
            ["Clustering methods", str(len(clustering.get("methods", {})))],
            ["Apriori rules", str(len(apriori.get("rules", [])))],
        ],
    )

    doc.add_heading("3. Startup and Execution Flow", level=1)
    add_numbered_list(
        doc,
        [
            "`start.sh` checks for Python, creates or reuses `.venv`, and installs dependencies only when `requirements.txt` changes.",
            "The script searches for a free local port if the default port is already occupied.",
            "If required preprocessing or JSON cache files are missing, the pipeline regenerates them before the dashboard opens.",
            "The visualization server then starts and serves the browser app from `user_tools/visualisation_tool/`.",
            "The browser opens automatically unless `OPEN_BROWSER=0` is set.",
        ],
    )

    doc.add_heading("4. Data Understanding and Preprocessing", level=1)
    doc.add_paragraph(
        "The preprocessing logic in `data_harness.py` follows the medical workflow expected by the course requirements and "
        "the reference papers. The goal is to produce a clinically interpretable, model-ready representation without "
        "leakage or meaningless categories."
    )
    add_bullet_list(
        doc,
        [
            "Keep the first encounter per patient to avoid repeated-visit leakage.",
            "Remove terminal discharges that cannot support a readmission prediction.",
            "Remove invalid gender values.",
            "Compute near-zero variance statistics for medication fields and drop the low-information ones.",
            "Replace missing values with explicit clinical labels such as `Unknown` and `Missing`.",
            "Convert age bands into numeric values for modeling.",
            "Group ICD-9 diagnosis codes into broad clinical categories.",
            "Create medication-change summaries before the medication columns are dropped.",
            "Build a final one-hot encoded model-ready matrix with standardized numeric features.",
        ],
    )

    doc.add_paragraph(
        "This stage produces the cleaned CSV, the model-ready CSV, and a family of JSON audit artifacts that the "
        "dashboard uses to show the user what was removed, why it was removed, and how the final feature space was built."
    )

    doc.add_heading("5. Predictive Analytics", level=1)
    doc.add_heading("5.1 Baseline Classification", level=2)
    doc.add_paragraph(
        "The baseline page compares multiple classifiers on the same stratified 80/20 split. The report includes "
        "metrics, ROC curves, confusion matrices, TN/FP/FN/TP counts, and a significance comparison between the best "
        "model and its closest competitor."
    )
    add_bullet_list(
        doc,
        [
            "Logistic Regression",
            "Random Forest",
            "Gradient Boosting",
            "XGBoost when available",
        ],
    )

    doc.add_heading("5.2 Feature Selection", level=2)
    doc.add_paragraph(
        "Feature selection is now performed on the full stratified training split. This avoids the earlier sample-based "
        "constraint and gives a stronger final ranking table."
    )
    add_bullet_list(
        doc,
        [
            "Mutual Information",
            "Chi-square",
            "Random Forest importance",
            "Absolute Logistic Regression coefficients",
            "RFE with Logistic Regression",
            "Consensus ranking across methods",
        ],
    )

    doc.add_heading("5.3 Feature-Subset Experiments", level=2)
    doc.add_paragraph(
        "The feature-subset page retrains Logistic Regression and Random Forest on top-ranked feature groups and compares "
        "ROC-AUC, accuracy, recall, and F1. This gives a compact view of how much predictive power is retained when the "
        "feature space is reduced."
    )

    doc.add_heading("5.4 Regression", level=2)
    doc.add_paragraph(
        "Regression is now fully integrated and predicts `time_in_hospital` using the model-ready matrix without the "
        "target leakage column. The report compares multiple regressors and records residual behavior."
    )
    add_bullet_list(
        doc,
        [
            "Linear Regression",
            "Ridge Regression",
            "Random Forest Regressor",
            "Gradient Boosting Regressor",
            "XGBoost Regressor when available",
        ],
    )

    doc.add_heading("6. Live Model Lab", level=1)
    doc.add_paragraph(
        "The live lab is the interactive part of the app. It allows the user to choose a classifier, tune hyperparameters, "
        "select a dataset size mode, and run the experiment directly from the browser through the local API in `server.py`."
    )
    add_bullet_list(
        doc,
        [
            "Auto safe",
            "Fast sample",
            "Medium sample",
            "Full split",
        ],
    )
    doc.add_paragraph(
        "The UI reports the selected data size, the train/test row counts, the resulting confusion matrix, the run status, "
        "and a history of recent experiments. This makes the app suitable for live presentation scenarios."
    )

    doc.add_heading("7. Descriptive Analytics", level=1)
    doc.add_paragraph(
        "The descriptive side of the app focuses on interpretable patterns rather than a single score. Clustering and Apriori "
        "are intentionally documented with dataset-size notes so the user can see what is fully computed and what is "
        "sampled for presentation practicality."
    )
    add_bullet_list(
        doc,
        [
            "K-Means on PCA coordinates with k from 2 to 8",
            "Hierarchical clustering on the same projection",
            "DBSCAN with multiple eps and min_samples combinations",
            "Apriori association mining from interpretable clinical fields",
        ],
    )
    doc.add_paragraph(
        "Clustering remains sample-based because large hierarchical runs and browser scatter rendering are not practical on "
        "the full 69,970-row model-ready table. Apriori is capped for the same reason: candidate growth becomes expensive "
        "on the full table."
    )

    doc.add_heading("8. Dashboard Pages and User Experience", level=1)
    add_bullet_list(
        doc,
        [
            "Raw data browser with feature-level drill-down",
            "Cleaning browser with removed vs retained rows",
            "Model-ready browser with encoded feature preview",
            "Baseline ML comparison page",
            "Regression evaluation page",
            "Live Model Lab with experiment reruns",
            "Clustering Lab with method and parameter controls",
            "Association Mining page with top rules and itemsets",
            "Conclusion page summarizing the strongest findings",
        ],
    )

    doc.add_heading("9. Known Constraints and Design Rationale", level=1)
    add_bullet_list(
        doc,
        [
            "Clustering and Apriori are intentionally sampled for responsiveness.",
            "The repo is dashboard-first rather than script-first, so standalone wrapper scripts would still improve submission clarity.",
            "The raw dataset is currently present in the repo and should be reviewed against the course rule before final submission.",
            "The dashboard is optimized for presentation and iterative testing, not for repeatedly rerunning the heaviest experiments from scratch on every launch.",
        ],
    )

    doc.add_heading("10. Final Interpretation", level=1)
    doc.add_paragraph(
        "The current application is a presentation-ready local analytics system. It opens quickly when cached artifacts exist, "
        "recomputes the pipeline when requested, and exposes the major course requirements in a coherent browser interface. "
        "Its strongest contribution is the combination of reproducible preprocessing, interactive predictive experiments, and "
        "transparent reporting of what was computed on full data versus what was sampled for practical reasons."
    )

    doc.add_paragraph()
    doc.add_paragraph("Reference artifacts used for this report are stored inside `user_tools/visualisation_tool/` and `data/`.")

    doc.save(OUT_PATH)
    print(f"Saved {OUT_PATH}")


if __name__ == "__main__":
    main()
